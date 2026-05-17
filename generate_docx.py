import docx

def create_model_docs():
    doc = docx.Document()
    doc.add_heading('Danh sách Bảng và Thuộc tính trong Hệ thống', 0)

    models = [
        {
            "name": "co_so (Cơ sở)",
            "attributes": [
                ("ma_co_so", "String(20), Primary Key"),
                ("ten_co_so", "String(100), Not Null"),
                ("dia_chi", "String(255)")
            ]
        },
        {
            "name": "khoa (Khoa)",
            "attributes": [
                ("ma_khoa", "String(20), Primary Key"),
                ("ten_khoa", "String(100), Not Null")
            ]
        },
        {
            "name": "sinh_vien (Sinh viên)",
            "attributes": [
                ("ma_sv", "String(20), Primary Key"),
                ("ho_ten", "String(100), Not Null"),
                ("ngay_sinh", "Date"),
                ("gioi_tinh", "String(10)"),
                ("ma_khoa", "String(20), Foreign Key (khoa.ma_khoa)"),
                ("ma_co_so", "String(20), Foreign Key (co_so.ma_co_so)"),
                ("da_xoa", "Boolean, Default=False")
            ]
        },
        {
            "name": "giang_vien (Giảng viên)",
            "attributes": [
                ("ma_gv", "String(20), Primary Key"),
                ("ho_ten", "String(100), Not Null"),
                ("hoc_vi", "String(50)"),
                ("ma_khoa", "String(20), Foreign Key (khoa.ma_khoa)"),
                ("ma_co_so", "String(20), Foreign Key (co_so.ma_co_so)"),
                ("da_xoa", "Boolean, Default=False")
            ]
        },
        {
            "name": "hoc_phan (Học phần)",
            "attributes": [
                ("ma_hp", "String(20), Primary Key"),
                ("ten_hp", "String(100), Not Null"),
                ("so_tin_chi", "Integer, Not Null"),
                ("ma_khoa", "String(20), Foreign Key (khoa.ma_khoa)")
            ]
        },
        {
            "name": "phong_hoc (Phòng học)",
            "attributes": [
                ("ma_phong", "String(20), Primary Key"),
                ("ten_phong", "String(100), Not Null"),
                ("suc_chua", "Integer"),
                ("ma_co_so", "String(20), Foreign Key (co_so.ma_co_so)")
            ]
        },
        {
            "name": "lop_hoc_phan (Lớp học phần)",
            "attributes": [
                ("ma_lop_hp", "String(20), Primary Key"),
                ("ma_hp", "String(20), Foreign Key (hoc_phan.ma_hp)"),
                ("ma_gv", "String(20)"),
                ("ma_phong", "String(20), Foreign Key (phong_hoc.ma_phong)"),
                ("ma_co_so", "String(20), Foreign Key (co_so.ma_co_so)"),
                ("hoc_ky", "Integer"),
                ("nam_hoc", "String(20)"),
                ("si_so_toi_da", "Integer"),
                ("so_luong_da_dang_ky", "Integer, Default=0"),
                ("version", "Integer, Default=0, Not Null (Optimistic Locking)")
            ]
        },
        {
            "name": "lich_hoc (Lịch học)",
            "attributes": [
                ("ma_lich", "String(20), Primary Key"),
                ("ma_lop_hp", "String(20), Foreign Key (lop_hoc_phan.ma_lop_hp)"),
                ("thu", "Integer"),
                ("tiet_bat_dau", "Integer"),
                ("tiet_ket_thuc", "Integer"),
                ("ma_phong", "String(20), Foreign Key (phong_hoc.ma_phong)"),
                ("ma_co_so", "String(20), Foreign Key (co_so.ma_co_so)")
            ]
        },
        {
            "name": "dang_ky (Đăng ký)",
            "attributes": [
                ("ma_dk", "Integer, Primary Key, Autoincrement"),
                ("ma_sv", "String(20), Not Null"),
                ("ma_lop_hp", "String(20), Foreign Key (lop_hoc_phan.ma_lop_hp)"),
                ("thoi_gian_dang_ky", "DateTime, Server Default=now()"),
                ("trang_thai", "String(50), Default='THANH_CONG'"),
                ("ma_co_so", "String(20), Foreign Key (co_so.ma_co_so)"),
                ("UNIQUE", "UNIQUE(ma_sv, ma_lop_hp)")
            ]
        }
    ]

    for model in models:
        doc.add_heading(model["name"], level=1)
        
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Tên thuộc tính'
        hdr_cells[1].text = 'Định kiểu & Ràng buộc'
        
        for attr in model["attributes"]:
            row_cells = table.add_row().cells
            row_cells[0].text = attr[0]
            row_cells[1].text = attr[1]
            
        doc.add_paragraph() # Add some space after the table

    doc.save('report/Danh_sach_bang_va_thuoc_tinh.docx')
    print("Document saved successfully!")

if __name__ == '__main__':
    create_model_docs()
